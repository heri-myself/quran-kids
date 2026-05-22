from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# --- Page margins ---
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3)
section.right_margin  = Cm(2.5)

# --- Styles helper ---
def heading1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
    p.runs[0].font.size = Pt(16)
    return p

def heading2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)
    p.runs[0].font.size = Pt(13)
    return p

def heading3(text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = RGBColor(0x37, 0x51, 0x8c)
    p.runs[0].font.size = Pt(11)
    return p

def body(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    run.bold = bold
    run.italic = italic
    return p

def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), 'EFF6FF')
    p._p.get_or_add_pPr().append(shading)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # Header row
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:fill'), 'DBEAFE')
        hdr[i].paragraphs[0]._p.get_or_add_pPr().append(shading)
    # Data rows
    for ri, row in enumerate(rows):
        cells = table.rows[ri+1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
            for run in cells[ci].paragraphs[0].runs:
                run.font.size = Pt(10)
    return table

# ══════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Dokumentasi Fine-Tuning Model Whisper')
run.font.size = Pt(22)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
run.font.name = 'Calibri'

doc.add_paragraph()

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = subtitle.add_run('untuk Tilawah Al-Quran Juz 30\nMenggunakan Dataset Al-Husary')
run2.font.size = Pt(14)
run2.font.color.rgb = RGBColor(0x37, 0x51, 0x8c)
run2.font.name = 'Calibri'

doc.add_paragraph()
doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run(f'Proyek  : Quran Kids App\n').font.size = Pt(11)
meta.add_run(f'Tanggal : {datetime.date.today().strftime("%d %B %Y")}\n').font.size = Pt(11)
meta.add_run(f'Base Model : naazimsnh02/whisper-large-v3-turbo-ar-quran').font.size = Pt(11)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 1. PENDAHULUAN
# ══════════════════════════════════════════════════════════
heading1('1. Pendahuluan')
body(
    'Dokumen ini menjelaskan proses fine-tuning model Whisper untuk meningkatkan '
    'akurasi transkripsi bacaan Al-Quran pada aplikasi Quran Kids. Fine-tuning '
    'dilakukan menggunakan audio murottal Mahmoud Khalil Al-Husary sebagai dataset '
    'referensi, khusus untuk Juz 30 (Surah 78–114).'
)

doc.add_paragraph()
heading2('1.1 Latar Belakang')
body(
    'Model Whisper yang digunakan saat ini (naazimsnh02/whisper-large-v3-turbo-ar-quran) '
    'sudah dioptimalkan untuk bahasa Arab umum, namun masih memiliki keterbatasan dalam '
    'mengenali bacaan Al-Quran dengan akurat, terutama untuk:'
)
bullet('Surat-surat pendek yang sering dibaca anak-anak (Juz 30)')
bullet('Variasi pelafalan yang berbeda dari teks Arab standar')
bullet('Ayat-ayat yang memiliki kemiripan bunyi dengan ayat lain')

doc.add_paragraph()
heading2('1.2 Tujuan')
bullet('Meningkatkan akurasi transkripsi ayat-ayat Juz 30')
bullet('Mengurangi kesalahan transkripsi pada surat-surat pendek seperti An-Nas, Al-Ikhlas, Al-Falaq')
bullet('Menghasilkan model yang siap dipakai di backend aplikasi')

doc.add_paragraph()

# ══════════════════════════════════════════════════════════
# 2. APA ITU FINE-TUNING
# ══════════════════════════════════════════════════════════
heading1('2. Apa itu Fine-Tuning?')
body(
    'Fine-tuning adalah proses melatih ulang model AI yang sudah ada menggunakan '
    'data yang lebih spesifik, sehingga model menjadi lebih akurat untuk domain tertentu.'
)

doc.add_paragraph()
heading2('2.1 Analogi Sederhana')
body(
    'Bayangkan Whisper seperti seorang sarjana bahasa Arab yang sudah lulus kuliah '
    '(trained dengan jutaan data). Fine-tuning ibarat mengirim sarjana tersebut ke '
    'pesantren selama beberapa minggu — dia sudah pintar, tinggal "diarahkan" untuk '
    'lebih fasih membaca Al-Quran secara spesifik.'
)

doc.add_paragraph()
heading2('2.2 Perbedaan Training dari Nol vs Fine-Tuning')
add_table(
    ['Aspek', 'Training dari Nol', 'Fine-Tuning'],
    [
        ['Data yang dibutuhkan', 'Jutaan sampel', 'Ratusan sampel'],
        ['Waktu komputasi', 'Berbulan-bulan', '45 menit – beberapa jam'],
        ['Biaya GPU', 'Sangat mahal', 'Gratis (Colab T4)'],
        ['Hasil', 'Model umum', 'Model spesifik domain'],
        ['Kualitas awal', 'Dimulai dari nol', 'Sudah pintar, tinggal diarahkan'],
    ]
)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════
# 3. DATASET
# ══════════════════════════════════════════════════════════
heading1('3. Dataset')
heading2('3.1 Qari Referensi: Mahmoud Khalil Al-Husary')
body(
    'Al-Husary dipilih sebagai dataset referensi karena:'
)
bullet('Diakui oleh Al-Azhar sebagai standar pengajaran tajwid dunia')
bullet('Pelafalan makhraj paling terpisah dan jelas — ideal untuk training model AI')
bullet('Tersedia versi Muallim (tempo lambat dengan jeda per ayat)')
bullet('Audio tersedia bebas di EveryAyah.com dengan kualitas 128kbps')

doc.add_paragraph()
heading2('3.2 Cakupan Dataset')
add_table(
    ['Parameter', 'Detail'],
    [
        ['Cakupan', 'Juz 30 — Surah 78 (An-Naba) hingga Surah 114 (An-Nas)'],
        ['Jumlah Surah', '37 surah'],
        ['Jumlah Ayat', '564 ayat'],
        ['Format Audio', 'MP3 128kbps → dikonversi ke WAV 16kHz mono'],
        ['Ukuran Total', '~110 MB (MP3) / ~180 MB (WAV)'],
        ['Sumber', 'https://everyayah.com/data/Husary_128kbps/'],
    ]
)

doc.add_paragraph()
heading2('3.3 Format Nama File Audio')
body('File audio menggunakan format: [3 digit surah][3 digit ayat].mp3')
code_block('Contoh:\n078001.mp3  → Surah An-Naba (78), Ayat 1\n114006.mp3  → Surah An-Nas (114), Ayat 6')

doc.add_paragraph()
heading2('3.4 Pembagian Dataset')
body('Dataset dibagi menjadi dua bagian secara acak (random seed=42):')
add_table(
    ['Bagian', 'Proporsi', 'Jumlah Ayat', 'Fungsi'],
    [
        ['Train', '90%', '~508 ayat', 'Data untuk melatih model'],
        ['Validation', '10%', '~56 ayat', 'Data untuk evaluasi setiap epoch'],
    ]
)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════
# 4. PROSES FINE-TUNING
# ══════════════════════════════════════════════════════════
heading1('4. Proses Fine-Tuning')
body('Fine-tuning terdiri dari 6 tahap utama yang berjalan secara berurutan:')

doc.add_paragraph()
heading2('Tahap 1 — Download Audio')
body(
    'Script otomatis mengunduh 564 file MP3 dari EveryAyah.com. '
    'File disimpan dengan nama terstruktur (078001.mp3 dst). '
    'Proses ini memakan waktu sekitar 15 menit.'
)
code_block(
    'URL format:\nhttps://everyayah.com/data/Husary_128kbps/[surah][ayat].mp3\n\n'
    'Contoh:\nhttps://everyayah.com/data/Husary_128kbps/114001.mp3'
)

doc.add_paragraph()
heading2('Tahap 2 — Pembuatan Teks Pasangan')
body(
    'Setiap file audio dipasangkan dengan teks Arab Uthmani-nya. '
    'Teks diambil dari API Quran.com dan dibersihkan dari nomor ayat Arab (١٢٣). '
    'Hasilnya berupa 564 pasangan: audio ↔ teks.'
)

doc.add_paragraph()
heading2('Tahap 3 — Preprocessing Audio')
body(
    'Semua file MP3 dikonversi ke format WAV dengan spesifikasi:'
)
bullet('Sample rate: 16.000 Hz (standar Whisper)')
bullet('Channel: Mono (1 channel)')
bullet('Format: 32-bit float PCM')
body('Konversi menggunakan library librosa + soundfile. Estimasi waktu: ~30 menit.')

doc.add_paragraph()
heading2('Tahap 4 — Preprocessing Model (Mel Spectrogram)')
body(
    'Audio diubah menjadi Mel Spectrogram — representasi visual suara yang '
    'bisa diproses oleh model neural network. Sumbu X adalah waktu, '
    'sumbu Y adalah frekuensi, dan warna menunjukkan intensitas suara. '
    'Teks Arab diubah menjadi token ID menggunakan tokenizer Whisper.'
)

doc.add_paragraph()
heading2('Tahap 5 — Training Loop')
body('Model dilatih dengan konfigurasi berikut:')
add_table(
    ['Parameter', 'Nilai', 'Keterangan'],
    [
        ['Epochs (maksimum)', '5', 'Bisa berhenti lebih awal jika early stopping aktif'],
        ['Batch size efektif', '16', 'Per-device 8, gradient accumulation 2'],
        ['Learning rate', '1e-5', 'Kecil agar tidak merusak pengetahuan sebelumnya'],
        ['Warmup steps', '50', 'Pemanasan awal sebelum learning rate penuh'],
        ['Precision', 'FP16', 'Half-precision untuk hemat memori GPU'],
        ['Early stopping', '2 epoch', 'Stop jika WER tidak turun selama 2 epoch berturut-turut'],
        ['Checkpoint', 'Setiap epoch', 'Model di-backup ke HuggingFace Hub otomatis'],
    ]
)

doc.add_paragraph()
body(
    'Setiap epoch, model dievaluasi menggunakan WER (Word Error Rate) pada '
    'data validasi. Model terbaik (WER terendah) disimpan dan digunakan sebagai output akhir.'
)

doc.add_paragraph()
heading2('Tahap 6 — Push ke HuggingFace Hub')
body(
    'Model terbaik diunggah otomatis ke HuggingFace Hub sebagai private repository. '
    'Dari sana, backend dapat mengunduh dan menggunakannya hanya dengan mengganti '
    'nama model di satu baris kode.'
)
code_block(
    '# Sebelum fine-tuning:\n'
    '"naazimsnh02/whisper-large-v3-turbo-ar-quran"\n\n'
    '# Setelah fine-tuning:\n'
    '"username/whisper-quran-juz30-husary"'
)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════
# 5. METRIK EVALUASI
# ══════════════════════════════════════════════════════════
heading1('5. Metrik Evaluasi (WER)')
body(
    'WER (Word Error Rate) adalah metrik standar untuk mengukur akurasi model ASR. '
    'Semakin kecil WER, semakin akurat model.'
)
code_block(
    'WER = (Substitusi + Penghapusan + Penyisipan) / Total Kata Referensi × 100%\n\n'
    'Contoh:\n'
    'Referensi : "قُلْ هُوَ اللَّهُ أَحَدٌ"\n'
    'Prediksi  : "قُلْ هُوَ اللَّهُ"\n'
    'WER       : 1/4 × 100 = 25%'
)
doc.add_paragraph()
add_table(
    ['WER', 'Interpretasi'],
    [
        ['0–5%', 'Sangat baik — hampir sempurna'],
        ['5–15%', 'Baik — layak produksi'],
        ['15–30%', 'Cukup — perlu lebih banyak data'],
        ['> 30%', 'Kurang — perlu perbaikan dataset atau konfigurasi'],
    ]
)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════
# 6. TOOLS & INFRASTRUKTUR
# ══════════════════════════════════════════════════════════
heading1('6. Tools dan Infrastruktur')

heading2('6.1 Google Colab')
body(
    'Training dijalankan di Google Colab menggunakan GPU NVIDIA T4 (gratis). '
    'Notebook dapat dijalankan dengan "Run All" dan ditinggalkan — '
    'proses berjalan otomatis di background hingga selesai (~45 menit untuk Juz 30).'
)

doc.add_paragraph()
heading2('6.2 Library Python')
add_table(
    ['Library', 'Versi', 'Fungsi'],
    [
        ['transformers', '≥4.36', 'Load, fine-tune, dan save model Whisper'],
        ['datasets', '≥2.14', 'Manajemen dataset HuggingFace'],
        ['accelerate', '≥0.24', 'Optimasi training multi-GPU/CPU'],
        ['evaluate + jiwer', 'latest', 'Hitung metrik WER'],
        ['librosa', '≥0.10', 'Load dan konversi audio'],
        ['soundfile', 'latest', 'Simpan file WAV'],
        ['torch + torchaudio', '≥2.0', 'Framework deep learning'],
    ]
)

doc.add_paragraph()
heading2('6.3 Estimasi Waktu dan Hardware')
add_table(
    ['Hardware', 'Dataset', 'Estimasi Waktu'],
    [
        ['Colab T4 (gratis)', 'Juz 30 (564 ayat)', '~45 menit'],
        ['Colab A100 (Pro)', 'Juz 30 (564 ayat)', '~15 menit'],
        ['Mac M-series (MPS)', 'Juz 30 (564 ayat)', '~3–4 jam'],
        ['Colab T4 (gratis)', 'Seluruh Quran (6.236 ayat)', '~2 hari*'],
        ['Colab A100 (Pro)', 'Seluruh Quran (6.236 ayat)', '~8 jam'],
    ]
)
body('*Colab gratis disconnect setiap ~3 jam. Checkpoint per-epoch memastikan progress tidak hilang.', italic=True)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════
# 7. CARA PAKAI HASIL
# ══════════════════════════════════════════════════════════
heading1('7. Cara Menggunakan Model Hasil Fine-Tuning')
heading2('7.1 Lokasi Model')
body(
    'Setelah training selesai, model tersimpan di HuggingFace Hub sebagai private repository. '
    'Model juga di-cache otomatis di Mac setelah pertama kali diunduh:'
)
code_block('~/.cache/huggingface/hub/models--username--whisper-quran-juz30-husary/')

doc.add_paragraph()
heading2('7.2 Integrasi ke Backend')
body('Cukup ganti satu baris di file tilawah_service.py:')
code_block(
    '# File: backend/python/tilawah_service.py\n\n'
    '# Baris lama:\n'
    'model = AutoModelForSpeechSeq2Seq.from_pretrained(\n'
    '    "naazimsnh02/whisper-large-v3-turbo-ar-quran", ...\n'
    ')\n\n'
    '# Ganti dengan:\n'
    'model = AutoModelForSpeechSeq2Seq.from_pretrained(\n'
    '    "username/whisper-quran-juz30-husary", ...\n'
    ')'
)

doc.add_paragraph()
heading2('7.3 Langkah Integrasi')
p = doc.add_paragraph(style='List Number')
p.add_run('Ganti nama model di tilawah_service.py')
p = doc.add_paragraph(style='List Number')
p.add_run('Restart backend Python')
p = doc.add_paragraph(style='List Number')
p.add_run('Backend otomatis download model (~800MB, sekali saja)')
p = doc.add_paragraph(style='List Number')
p.add_run('Model ter-cache, tidak perlu download ulang di restart berikutnya')
p = doc.add_paragraph(style='List Number')
p.add_run('Test dengan membaca beberapa ayat Juz 30 di aplikasi')

doc.add_paragraph()

# ══════════════════════════════════════════════════════════
# 8. RENCANA PENGEMBANGAN
# ══════════════════════════════════════════════════════════
heading1('8. Rencana Pengembangan Selanjutnya')
add_table(
    ['Fase', 'Fitur', 'Estimasi Waktu'],
    [
        ['Fase 1 (selesai)', 'Fine-tuning Juz 30 dengan Al-Husary', '1 minggu'],
        ['Fase 2', 'Expand ke Juz 29–28 (tambah data)', '1 minggu'],
        ['Fase 3', 'WhisperX forced alignment → deteksi mad', '1 minggu'],
        ['Fase 4', 'Bandingkan user vs referensi Al-Husary (ritme)', '2 minggu'],
        ['Fase 5', 'Fine-tuning seluruh 30 Juz', '1 bulan'],
    ]
)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════
# 9. REFERENSI
# ══════════════════════════════════════════════════════════
heading1('9. Referensi')
bullet('HuggingFace Fine-Tuning Whisper Guide: https://huggingface.co/blog/fine-tune-whisper')
bullet('Model Base: https://huggingface.co/naazimsnh02/whisper-large-v3-turbo-ar-quran')
bullet('Dataset Audio: https://everyayah.com/data/Husary_128kbps/')
bullet('Teks Quran API: https://api.qurancdn.com')
bullet('Notebook Colab: backend/python/finetune_whisper_juz30.ipynb')

# Save
output_path = '/Applications/Works/QuranBest/Vibe Code/Quran Kids/docs/Dokumentasi_FineTuning_Whisper_Juz30.docx'
doc.save(output_path)
print(f'✅ Dokumen disimpan: {output_path}')
